"""Extract plain text from uploaded documents: PDF, DOCX, XLSX, images, text."""
from __future__ import annotations

import base64
import io
import os
from dataclasses import dataclass

from ..config import get_settings

_IMAGE_MEDIA = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}
_TEXT_EXT = {".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log", ".rst"}
_MAX_XLSX_ROWS = 5000


@dataclass
class ExtractedDoc:
    title: str
    text: str


def supported_extensions() -> list[str]:
    return sorted(
        {".pdf", ".docx", ".xlsx", ".xls", *_TEXT_EXT, *_IMAGE_MEDIA.keys()}
    )


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _base_title(filename: str) -> str:
    return os.path.basename(filename or "document")


# ---------------------------------------------------------------- per-format
def _page_slice(total: int, start_page: int | None, end_page: int | None) -> tuple[int, int]:
    """Convert 1-based inclusive (start_page, end_page) to a 0-based [s, e) range."""
    s = 0 if not start_page or start_page < 1 else start_page - 1
    s = min(s, total)
    e = total if not end_page or end_page < 1 else min(end_page, total)
    if e < s:
        e = s
    return s, e


def _ocr_pdf_pages(data: bytes, page_indices: list[int], filename: str) -> tuple[list[str], str | None]:
    """Render scanned PDF pages to images and transcribe them via the vision model.

    Returns (texts, error). `error` is set when nothing could be transcribed so the
    caller can surface a real reason instead of a generic "no text" message.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        try:
            import pymupdf as fitz  # newer import name
        except ImportError:
            return [], (
                "This PDF has no text layer and OCR is unavailable: PyMuPDF is not installed. "
                "Run `pip install -r requirements.txt` (or `pip install pymupdf`) and restart the backend."
            )

    settings = get_settings()
    indices = page_indices[: settings.pdf_ocr_max_pages]
    out: list[str] = []
    first_error: str | None = None
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # noqa: BLE001
        return [], f"Could not open PDF for OCR: {exc}"

    try:
        for i in indices:
            try:
                pix = doc[i].get_pixmap(dpi=settings.pdf_ocr_dpi)
                png = pix.tobytes("png")
                txt = _extract_image(png, "image/png", f"{filename} (page {i + 1})").text
            except Exception as exc:  # noqa: BLE001
                first_error = first_error or str(exc)
                txt = ""
            if txt.strip():
                out.append(txt.strip())
    finally:
        doc.close()

    if out:
        return out, None
    return out, first_error or "OCR produced no text from the rendered pages."


def _extract_pdf(
    data: bytes, filename: str, start_page: int | None = None, end_page: int | None = None
) -> ExtractedDoc:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    total = len(reader.pages)
    s, e = _page_slice(total, start_page, end_page)

    parts: list[str] = []
    for i in range(s, e):
        try:
            txt = reader.pages[i].extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            parts.append(txt.strip())

    # Scanned / image-only PDF: pypdf finds (almost) no text → OCR the page range.
    if sum(len(p) for p in parts) < 100:
        ocr_parts, ocr_error = _ocr_pdf_pages(data, list(range(s, e)), filename)
        if ocr_parts:
            parts = ocr_parts
        elif ocr_error:
            # Propagate the real reason; the ingest service reports it to the UI.
            raise RuntimeError(ocr_error)
        else:
            parts = []

    title = _base_title(filename)
    try:
        meta_title = (reader.metadata or {}).get("/Title")  # type: ignore[union-attr]
        if meta_title:
            title = str(meta_title)
    except Exception:
        pass
    return ExtractedDoc(title=title, text="\n\n".join(parts))


def _extract_docx(data: bytes, filename: str) -> ExtractedDoc:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return ExtractedDoc(title=_base_title(filename), text="\n\n".join(parts))


def _extract_xlsx(data: bytes, filename: str) -> ExtractedDoc:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
            count += 1
            if count >= _MAX_XLSX_ROWS:
                parts.append("… (truncated)")
                break
    wb.close()
    return ExtractedDoc(title=_base_title(filename), text="\n".join(parts))


def _extract_text(data: bytes, filename: str) -> ExtractedDoc:
    text = data.decode("utf-8", errors="replace")
    return ExtractedDoc(title=_base_title(filename), text=text)


def _extract_image(data: bytes, media_type: str, filename: str) -> ExtractedDoc:
    """Use the OpenAI vision model to transcribe text and describe the image."""
    from ..llm.client import get_client

    settings = get_settings()
    b64 = base64.b64encode(data).decode("ascii")
    resp = get_client().chat.completions.create(
        model=settings.llm_model,
        max_tokens=1500,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Transcribe ALL text visible in this image verbatim. "
                        "Then add a short factual description of any diagrams, charts, "
                        "tables, or notable visual content. Do not invent details.",
                    },
                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{b64}"}},
                ],
            }
        ],
    )
    text = resp.choices[0].message.content or ""
    return ExtractedDoc(title=_base_title(filename), text=text)


# ------------------------------------------------------------------ dispatch
def extract_file(
    filename: str,
    content_type: str | None,
    data: bytes,
    *,
    start_page: int | None = None,
    end_page: int | None = None,
) -> ExtractedDoc:
    """Extract text. start_page/end_page (1-based, inclusive) bound page-based
    formats (PDF); they are ignored for formats without a page concept."""
    ext = _ext(filename)
    if ext == ".pdf":
        return _extract_pdf(data, filename, start_page, end_page)
    if ext == ".docx":
        return _extract_docx(data, filename)
    if ext in (".xlsx", ".xls"):
        return _extract_xlsx(data, filename)
    if ext in _IMAGE_MEDIA:
        return _extract_image(data, _IMAGE_MEDIA[ext], filename)
    if ext in _TEXT_EXT:
        return _extract_text(data, filename)
    # Fall back on content-type sniffing, else best-effort text decode.
    if content_type:
        if content_type == "application/pdf":
            return _extract_pdf(data, filename, start_page, end_page)
        if content_type.startswith("image/"):
            return _extract_image(data, content_type, filename)
        if content_type.startswith("text/"):
            return _extract_text(data, filename)
    return _extract_text(data, filename)
